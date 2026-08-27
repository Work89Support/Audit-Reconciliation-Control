from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/คู่มือการตั้งชื่อไฟล์สำหรับส่งตรวจ_Audit_AI.docx")

BLUE = "0B67D0"
DARK = "10233F"
MUTED = "66758A"
LIGHT_BLUE = "EAF3FF"
LIGHT_GRAY = "F5F7FA"
GREEN = "16853B"
LIGHT_GREEN = "EAF7EE"
RED = "B42318"
LIGHT_RED = "FFF0EE"
GOLD = "A15C00"
LIGHT_GOLD = "FFF7E6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=DARK, italic=False, font="Arial Unicode MS"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=6, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_para(doc, text="", size=11, bold=False, color=DARK, before=0, after=6,
             align=None, italic=False, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, before, after)
    if align is not None:
        p.alignment = align
    p.paragraph_format.keep_with_next = keep
    set_run(p.add_run(text), size, bold, color, italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(p, after=4)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, after=4)
    p.add_run(text)
    return p


def add_callout(doc, label, body, fill=LIGHT_BLUE, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=2)
    set_run(p.add_run(label + "  "), 11, True, color)
    set_run(p.add_run(body), 11, False, DARK)
    add_para(doc, "", after=2)
    return table


def add_code_line(doc, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.15)
    set_run(p.add_run(text), 10.5, True, DARK, font="Courier New")
    return table


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line=1.15)
        set_run(p.add_run(header), font_size, True, DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            style_paragraph(p, after=0, line=1.15)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), font_size, False, DARK)
    set_table_geometry(table, widths)
    add_para(doc, "", after=2)
    return table


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("หน้า ")
    set_run(run, 9, False, MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        st = styles[name]
        st.font.name = "Arial Unicode MS"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]
        st.font.name = "Arial Unicode MS"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    configure_styles(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(hp.add_run("AUDIT AI  |  FILE INTAKE GUIDE"), 8.5, True, MUTED)
    set_page_number(section.footer.paragraphs[0])

    # Customer-pack opening block using compact_reference_guide tokens.
    add_para(doc, "AUDIT AI", 11, True, BLUE, after=5)
    add_para(doc, "คู่มือการตั้งชื่อไฟล์สำหรับส่งตรวจ", 26, True, DARK, after=4)
    add_para(doc, "มาตรฐานสำหรับทีมผู้ส่งไฟล์ เพื่อให้ระบบจำแนก อ่าน และกระทบยอดได้ครบถ้วน", 13, False, MUTED, after=16)
    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [4680, 4680])
    values = [
        ("ใช้กับ", "ไฟล์ STM / BO / PM / ฝากมือ / ถอน / หลักฐาน"),
        ("วันที่ในชื่อไฟล์", "YYYY-MM-DD เช่น 2026-08-24"),
    ]
    for r, (label, value) in enumerate(values):
        set_cell_shading(meta.cell(r, 0), LIGHT_BLUE)
        for c, text in enumerate((label, value)):
            p = meta.cell(r, c).paragraphs[0]
            style_paragraph(p, after=0)
            set_run(p.add_run(text), 10, c == 0, BLUE if c == 0 else DARK)
    add_para(doc, "", after=4)
    add_callout(doc, "กฎสำคัญที่สุด", "ชื่อไฟล์ทุกไฟล์ต้องมี บริษัท + ประเภท + วันที่ ห้ามใส่ข้อมูลเหล่านี้ไว้เฉพาะหัวข้ออีเมล", LIGHT_GOLD, GOLD)

    add_heading(doc, "1. รูปแบบชื่อไฟล์มาตรฐาน", 1)
    add_para(doc, "ตั้งชื่อโดยเรียงข้อมูลตามสูตรนี้ และใช้เครื่องหมายขีดล่าง (_) คั่นแต่ละส่วน", after=5)
    add_code_line(doc, "COMPANY_TYPE_SOURCE_NAME_DIRECTION_YYYY-MM-DD.ext")
    add_para(doc, "SOURCE คือธนาคารหรือ Provider, NAME คือชื่อเจ้าของบัญชี (ใส่เฉพาะไฟล์ธนาคาร) และ DIRECTION ใช้รหัสสั้น D / W / DW หากมีหลายไฟล์ให้เติม _01, _02 ต่อท้ายวันที่", color=MUTED, after=8)

    add_heading(doc, "2. รหัสบริษัทที่อนุญาต", 1)
    add_table(doc, ["รหัส", "บริษัท/ระบบที่ใช้"], [
        ("3XB", "ระบบ XXX"), ("AT4", "ระบบ 123"), ("FR8", "ระบบ 123"),
        ("MC8", "ระบบ XXX"), ("MR9", "ระบบ XXX"), ("PS8", "ระบบ XXX"),
        ("SK8", "ระบบ 123"), ("UFABET7M", "ระบบ UFABET7M"), ("UR9", "ระบบ XXX"),
    ], [1800, 7560])
    add_callout(doc, "ระวังชื่อย่อ", "ให้ใช้ UFABET7M เท่านั้น ไม่ใช้ UFA7M หรือ U7M และใช้ MC8 ไม่ใช้ MC เพื่อไม่ให้ระบบสร้างบริษัทซ้ำ", LIGHT_RED, RED)

    add_heading(doc, "3. รหัสประเภทไฟล์และนิยาม", 1)
    type_rows = [
        ("STM", "Statement หรือรายการเดินบัญชีธนาคาร", "PDF/XLSX/CSV"),
        ("BO", "รายงานหน้า BO หรือรายงานหลังบ้านของบริษัท", "XLSX/CSV"),
        ("PM", "รายการฝากหรือถอนจาก Payment Provider", "XLSX/CSV"),
        ("MANUAL_CREDIT", "รายการเติมเครดิตด้วยมือ", "XLSX/CSV"),
        ("MANUAL_PAYMENT", "รายการฝากมือฝั่ง Payment", "XLSX/CSV"),
        ("MANUAL_BONUS", "รายการโบนัสที่ทำด้วยมือ", "XLSX/CSV"),
        ("COMMISSION_WITHDRAW", "รายการถอนค่าคอมมิชชั่น", "XLSX/CSV"),
        ("CREDIT_WITHDRAW", "รายงานถอนเครดิต", "XLSX/CSV"),
        ("EVIDENCE", "ไฟล์ชี้แจงหรือหลักฐานเพิ่มเติม", "PDF/DOCX/XLSX"),
    ]
    add_table(doc, ["รหัส", "ใช้เมื่อ", "ไฟล์ที่แนะนำ"], type_rows, [2520, 5100, 1740], font_size=9)
    add_callout(doc, "ไม่ควรส่งเป็น UNKNOWN", "“ยังไม่ทราบประเภท” ใช้เฉพาะภายในระบบตอนรอตรวจ ไม่ควรใช้เป็นชื่อไฟล์ที่ทีมส่งมา", LIGHT_GOLD, GOLD)

    add_heading(doc, "4. รูปแบบไฟล์ธนาคารและรหัสฝาก–ถอนแบบสั้น", 1)
    add_para(doc, "ไฟล์ธนาคารต้องมีชื่อเจ้าของบัญชี เพื่อแยกบัญชีที่ใช้กระทบยอดได้ถูกต้อง", after=5)
    add_code_line(doc, "COMPANY_STM_BANK_ACCOUNTNAME_D/W/DW_YYYY-MM-DD.ext", fill=LIGHT_GREEN)
    add_table(doc, ["รหัส", "ความหมาย", "ตัวอย่าง"], [
        ("D", "ฝาก (Deposit)", "UFABET7M_STM_KB_เพ็ญศรี_D_2026-08-10.pdf"),
        ("W", "ถอน (Withdraw)", "UFABET7M_STM_KB_เพ็ญศรี_W_2026-08-10.pdf"),
        ("DW", "มีทั้งฝากและถอน", "UFABET7M_STM_KB_เพ็ญศรี_DW_2026-08-10.pdf"),
    ], [1200, 2500, 5660], font_size=9)
    add_callout(doc, "ตัวอย่างจากทีม", "รายการถอน_KB_เพ็ญศรี_10_08_69.pdf ให้เปลี่ยนเป็น UFABET7M_STM_KB_เพ็ญศรี_W_2026-08-10.pdf", LIGHT_BLUE, BLUE)
    add_para(doc, "รหัสธนาคารที่แนะนำ", bold=True, after=3)
    add_table(doc, ["รหัส", "ธนาคาร", "รหัส", "ธนาคาร"], [
        ("KB", "กสิกรไทย", "SCB", "ไทยพาณิชย์"),
        ("KTB", "กรุงไทย", "BBL", "กรุงเทพ"),
        ("BAY", "กรุงศรีอยุธยา", "TTB", "ทีเอ็มบีธนชาต"),
        ("GSB", "ออมสิน", "UOB", "ยูโอบี"),
    ], [1200, 3480, 1200, 3480], font_size=9)
    add_para(doc, "ไฟล์ PM ใช้สูตรสั้น โดยไม่ต้องใส่ชื่อเจ้าของบัญชี", bold=True, before=5, after=3)
    add_code_line(doc, "COMPANY_PM_PROVIDER_D/W/DW_YYYY-MM-DD.ext")
    add_table(doc, ["หมวด", "ค่าที่ให้ใช้", "ตัวอย่าง"], [
        ("Provider", "AUTOPEER / AZPAY / CYBERPLUS / MYPAY / 12PAY", "AZPAY"),
        ("ทิศทาง", "D = ฝาก", "AT4_PM_AZPAY_D_2026-08-24.xlsx"),
        ("ทิศทาง", "W = ถอน", "AT4_PM_AZPAY_W_2026-08-24.xlsx"),
        ("ทิศทาง", "DW = มีทั้งฝากและถอน", "AT4_PM_AZPAY_DW_2026-08-24.xlsx"),
    ], [1600, 4100, 3660], font_size=9)

    add_heading(doc, "5. ตัวอย่างชื่อไฟล์ที่ถูกต้อง", 1)
    examples = [
        ("STM ธนาคารฝาก", "3XB_STM_SCB_สมชาย_D_2026-08-24.pdf"),
        ("STM ธนาคารถอน", "UFABET7M_STM_KB_เพ็ญศรี_W_2026-08-10.pdf"),
        ("STM ฝากและถอน", "FR8_STM_BBL_กิตติ_DW_2026-08-24.pdf"),
        ("BO หลังบ้าน", "FR8_BO_2026-08-24.xlsx"),
        ("PM ฝาก", "AT4_PM_AUTOPEER_D_2026-08-24.xlsx"),
        ("PM ถอน", "MC8_PM_AZPAY_W_2026-08-24.csv"),
        ("PM ฝากและถอน", "SK8_PM_CYBERPLUS_DW_2026-08-24.xlsx"),
        ("ฝากมือเครดิต", "PS8_MANUAL_CREDIT_2026-08-24.xlsx"),
        ("ฝากมือ Payment", "MR9_MANUAL_PAYMENT_2026-08-24.xlsx"),
        ("ฝากมือโบนัส", "UR9_MANUAL_BONUS_2026-08-24.xlsx"),
        ("ถอนค่าคอม", "3XB_COMMISSION_WITHDRAW_2026-08-24.xlsx"),
        ("ถอนเครดิต", "FR8_CREDIT_WITHDRAW_2026-08-24.xlsx"),
        ("หลักฐาน", "UFABET7M_EVIDENCE_CYBERPLUS_2026-08-24.docx"),
        ("หลายไฟล์", "UFABET7M_STM_KB_เพ็ญศรี_W_2026-08-24_02.pdf"),
    ]
    add_table(doc, ["กรณี", "ชื่อไฟล์"], examples, [2400, 6960])

    add_heading(doc, "6. ตัวอย่างที่ไม่ควรใช้ และวิธีแก้", 1)
    bad_rows = [
        ("รายงานล่าสุด.xlsx", "ไม่ทราบบริษัท ประเภท และวันที่", "3XB_BO_2026-08-24.xlsx"),
        ("ฝากถอน.pdf", "ไม่ทราบบริษัท ธนาคาร ชื่อบัญชี และวันที่", "FR8_STM_KB_สมชาย_DW_2026-08-24.pdf"),
        ("mc.xlsx", "ชื่อบริษัทย่อผิดและไม่ทราบประเภท", "MC8_BO_2026-08-24.xlsx"),
        ("U7M เติมมือ.docx", "ชื่อบริษัทไม่ตรงมาตรฐาน", "UFABET7M_EVIDENCE_2026-08-24.docx"),
        ("AZPAY.csv", "ไม่ทราบบริษัท ฝาก/ถอน และวันที่", "MR9_PM_AZPAY_D_2026-08-24.csv"),
        ("แก้ไขล่าสุดจริงสุด.xlsx", "เสี่ยงซ้ำและไม่รู้ลำดับแก้ไข", "AT4_BO_2026-08-24_REV01.xlsx"),
    ]
    add_table(doc, ["ชื่อที่ไม่ควรใช้", "ปัญหา", "แก้เป็น"], bad_rows, [2500, 3100, 3760], header_fill=LIGHT_RED, font_size=8.7)

    add_heading(doc, "7. เทมเพลตหัวข้ออีเมล", 1)
    add_para(doc, "หัวข้อเมลทั่วไป", bold=True, after=3)
    add_code_line(doc, "[AUDIT] COMPANY | YYYY-MM-DD | FILE TYPES")
    add_para(doc, "ตัวอย่าง", bold=True, after=3, before=6)
    add_code_line(doc, "[AUDIT] 3XB | 2026-08-24 | STM + BO + PM", fill=LIGHT_GREEN)
    add_para(doc, "หัวข้อเมลสำหรับไฟล์ชี้แจง", bold=True, after=3, before=6)
    add_code_line(doc, "[AUDIT-EVIDENCE] 3XB | 2026-08-24 | AZPAY")
    add_callout(doc, "หัวข้อเมลช่วยคัดแยกเมล", "แต่ระบบยังอ่านไฟล์แต่ละไฟล์แยกกัน จึงต้องตั้งชื่อไฟล์ให้ครบทุกไฟล์เสมอ", LIGHT_BLUE, BLUE)

    add_heading(doc, "8. วิธีส่งไฟล์ทีละขั้น", 1)
    steps = [
        "ตรวจว่ารหัสบริษัทตรงกับรายการที่อนุญาต 9 บริษัท",
        "ระบุว่าไฟล์เป็น STM, BO, PM, รายการฝากมือ/ถอน หรือ EVIDENCE",
        "ถ้าเป็น STM ธนาคาร ให้ใส่รหัสธนาคาร ชื่อเจ้าของบัญชี และ D/W/DW",
        "ถ้าเป็น PM ให้ใส่ชื่อ Provider และ D/W/DW",
        "เปลี่ยนวันที่เป็น YYYY-MM-DD เช่น 2026-08-24",
        "ตั้งชื่อไฟล์ตามสูตรและเปิดไฟล์ตรวจว่าไม่เสียหรือใส่รหัสผ่าน",
        "ตั้งหัวข้ออีเมลตามเทมเพลต แล้วแนบไฟล์ทั้งหมด",
        "ก่อนส่ง ตรวจรายการตามเช็กลิสต์ด้านล่างอีกครั้ง",
    ]
    for step in steps:
        add_number(doc, step)

    add_heading(doc, "9. เช็กลิสต์ก่อนกดส่ง", 1)
    checks = [
        "☐ ชื่อไฟล์มีรหัสบริษัทที่ถูกต้อง",
        "☐ ชื่อไฟล์มีประเภท เช่น STM / BO / PM / EVIDENCE",
        "☐ วันที่เป็น YYYY-MM-DD และตรงกับข้อมูลภายในไฟล์",
        "☐ ไฟล์ STM มีรหัสธนาคาร ชื่อเจ้าของบัญชี และ D/W/DW",
        "☐ ไฟล์ PM มี Provider และรหัส D/W/DW",
        "☐ ไฟล์เปิดได้ ไม่เสีย และไม่ติดรหัสผ่าน",
        "☐ ไม่มีชื่อไฟล์ซ้ำ หากซ้ำเติม _01, _02",
        "☐ ไฟล์แก้ไขใช้ REV01, REV02 และไม่ใช้คำว่า “ล่าสุด”",
        "☐ หัวข้ออีเมลมีบริษัท วันที่ และประเภทไฟล์",
    ]
    for item in checks:
        add_bullet(doc, item)

    add_heading(doc, "10. การส่งไฟล์แก้ไข", 1)
    add_para(doc, "ห้ามใช้คำว่า “แก้ไขล่าสุด” หรือเขียนทับชื่อเดิมโดยไม่ระบุรุ่น ให้เติม REV ตามลำดับ", after=5)
    add_code_line(doc, "AT4_BO_2026-08-24_REV01.xlsx")
    add_code_line(doc, "AT4_BO_2026-08-24_REV02.xlsx", fill=LIGHT_GREEN)
    add_callout(doc, "หลังส่งไฟล์แก้ไข", "ระบบจะเก็บไฟล์เดิมไว้เป็นประวัติ และใช้รุ่นล่าสุดที่ผ่านการตรวจเพื่อรันกระทบยอดใหม่", LIGHT_GREEN, GREEN)

    add_heading(doc, "11. ข้อความพร้อมส่งให้ทีม", 1)
    message = (
        "กรุณาตั้งชื่อไฟล์ก่อนส่งตามรูปแบบ:\n"
        "COMPANY_TYPE_SOURCE_NAME_DIRECTION_YYYY-MM-DD.ext\n\n"
        "ธนาคาร: UFABET7M_STM_KB_เพ็ญศรี_W_2026-08-10.pdf\n"
        "PM: AT4_PM_AUTOPEER_D_2026-08-24.xlsx\n"
        "รหัสบริษัท: 3XB, AT4, FR8, MC8, MR9, PS8, SK8, UFABET7M, UR9\n"
        "ประเภทหลัก: STM, BO, PM, MANUAL_CREDIT, MANUAL_PAYMENT, MANUAL_BONUS, "
        "COMMISSION_WITHDRAW, CREDIT_WITHDRAW, EVIDENCE\n\n"
        "รหัสฝากถอน: D = ฝาก, W = ถอน, DW = ฝากและถอน\n\n"
        "กรุณาใช้วันที่ YYYY-MM-DD และต้องใส่บริษัท/ประเภทไว้ในชื่อไฟล์ทุกไฟล์ "
        "หากมีหลายไฟล์ให้เติม _01, _02 และหากเป็นไฟล์แก้ไขให้เติม REV01, REV02"
    )
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.25)
    for idx, line in enumerate(message.split("\n")):
        if idx:
            p.add_run().add_break()
        set_run(p.add_run(line), 10, line.startswith("COMPANY_") or line.startswith("ตัวอย่าง:"), DARK)

    add_heading(doc, "สรุปแบบสั้น", 1)
    add_callout(doc, "ชื่อไฟล์ที่ดีต้องตอบได้ 5 ข้อ", "ของบริษัทอะไร · เป็นไฟล์อะไร · ช่องทางใด · ชื่อบัญชีอะไร (กรณีธนาคาร) · ฝาก/ถอนและวันที่ใด", LIGHT_GREEN, GREEN)
    add_para(doc, "เวอร์ชันเอกสาร: 1.1 | จัดทำสำหรับระบบ Audit AI Reconciliation Control", 9, False, MUTED, before=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
